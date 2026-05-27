"""
Document Format Agent v2.0 — 分层文档格式化引擎
支持: 大标题 / 副标题 / 小标题 / 正文 / 页面边距 独立设置
用法:
    python doc_format_agent.py --stdin --quiet   (API 格式化模式, 从 stdin 读 JSON)
    python doc_format_agent.py --stdin --preview  (预览模式, 输出文档结构)
    python doc_format_agent.py --stdin --text     (文本模式, 从文字创建 docx)
    python doc_format_agent.py --interactive      (交互模式)
"""

import argparse, sys, os, json
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8", errors="replace")
sys.stdin.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


FONT_LIST = ["宋体", "黑体", "微软雅黑", "仿宋", "楷体", "Arial", "Times New Roman"]
SIZE_LIST = ["10", "10.5", "11", "12", "14", "16", "18", "22", "24", "26"]
SPACING_LIST = ["1.0", "1.15", "1.5", "2.0", "2.5"]

DEFAULT_ELEMENTS = {
    "title":    {"font": "黑体", "size": 22, "bold": True,  "align": "center",  "space_before": 0,  "space_after": 12, "color": "000000", "line_spacing": 1.5},
    "subtitle": {"font": "黑体", "size": 16, "bold": False, "align": "center",  "space_before": 0,  "space_after": 6,  "color": "333333", "line_spacing": 1.5},
    "heading":  {"font": "黑体", "size": 14, "bold": True,  "align": "left",    "space_before": 12, "space_after": 6,  "color": "000000", "line_spacing": 1.5},
    "body":     {"font": "宋体", "size": 12, "bold": False, "align": "justify", "space_before": 0,  "space_after": 0,  "color": "000000", "line_spacing": 1.5, "first_line_indent": 0.74},
    "page":     {"margin_top": 2.5, "margin_bottom": 2.5, "margin_left": 3.0, "margin_right": 3.0},
}

PRESETS = {
    "学术论文": {
        "title": {"font": "黑体", "size": 22, "bold": True,  "align": "center", "space_before": 0, "space_after": 12, "color": "000000", "line_spacing": 1.5},
        "subtitle": {"font": "黑体", "size": 15, "bold": False, "align": "center", "space_before": 0, "space_after": 6, "color": "000000", "line_spacing": 1.5},
        "heading":  {"font": "黑体", "size": 14, "bold": True,  "align": "left", "space_before": 12, "space_after": 6, "color": "000000", "line_spacing": 1.5},
        "body":     {"font": "宋体", "size": 12, "bold": False, "align": "justify", "space_before": 0, "space_after": 0, "color": "000000", "line_spacing": 1.5, "first_line_indent": 0.74},
        "page":     {"margin_top": 2.5, "margin_bottom": 2.5, "margin_left": 3.0, "margin_right": 3.0},
    },
    "商业报告": {
        "title": {"font": "微软雅黑", "size": 26, "bold": True,  "align": "center", "space_before": 24, "space_after": 12, "color": "1a1a1a", "line_spacing": 1.15},
        "subtitle": {"font": "微软雅黑", "size": 16, "bold": False, "align": "center", "space_before": 0, "space_after": 18, "color": "666666", "line_spacing": 1.15},
        "heading":  {"font": "微软雅黑", "size": 14, "bold": True,  "align": "left", "space_before": 18, "space_after": 6, "color": "1a1a1a", "line_spacing": 1.15},
        "body":     {"font": "微软雅黑", "size": 11, "bold": False, "align": "left", "space_before": 0, "space_after": 6, "color": "333333", "line_spacing": 1.15, "first_line_indent": 0},
        "page":     {"margin_top": 2.0, "margin_bottom": 2.0, "margin_left": 2.5, "margin_right": 2.5},
    },
    "政府公文": {
        "title": {"font": "方正小标宋简体", "size": 22, "bold": True,  "align": "center", "space_before": 0, "space_after": 0, "color": "000000", "line_spacing": 1.5},
        "subtitle": {"font": "楷体", "size": 16, "bold": False, "align": "center", "space_before": 0, "space_after": 0, "color": "000000", "line_spacing": 1.5},
        "heading":  {"font": "黑体", "size": 16, "bold": True,  "align": "left", "space_before": 0, "space_after": 0, "color": "000000", "line_spacing": 1.5},
        "body":     {"font": "仿宋", "size": 16, "bold": False, "align": "justify", "space_before": 0, "space_after": 0, "color": "000000", "line_spacing": 1.5, "first_line_indent": 0.74},
        "page":     {"margin_top": 3.7, "margin_bottom": 3.5, "margin_left": 2.8, "margin_right": 2.6},
    },
    "小说/散文": {
        "title": {"font": "楷体", "size": 22, "bold": True,  "align": "center", "space_before": 0, "space_after": 12, "color": "000000", "line_spacing": 2.0},
        "subtitle": {"font": "楷体", "size": 15, "bold": False, "align": "center", "space_before": 0, "space_after": 18, "color": "444444", "line_spacing": 2.0},
        "heading":  {"font": "楷体", "size": 14, "bold": True,  "align": "left", "space_before": 24, "space_after": 6, "color": "000000", "line_spacing": 2.0},
        "body":     {"font": "楷体", "size": 14, "bold": False, "align": "justify", "space_before": 0, "space_after": 0, "color": "000000", "line_spacing": 2.0, "first_line_indent": 0.74},
        "page":     {"margin_top": 2.0, "margin_bottom": 2.0, "margin_left": 2.5, "margin_right": 2.5},
    },
}


class DocumentFormatter:
    """分层文档格式化器"""

    def __init__(self, config: dict):
        self.elements = config.get("elements", {})
        self.page = config.get("page", {})
        self.style_map = config.get("style_map", {})
        self.text_structure = config.get("text_structure", {})
        self._structure = None  # 智能结构检测结果缓存

    # ── 智能文档结构检测 ─────────────────────────────────

    def _detect_structure(self, doc) -> list:
        """扫描全文，智能识别每个段落的类型和层级

        综合判断信号：
          1. Word 样式名 (Heading 1/2/3, 标题 1/2/3)
          2. 编号模式 (第X章 / 一、 / 1.1 / 1. 等)
          3. 字号大小（相对正文的偏移）
          4. 文本长度 + 结束标点
          5. 是否加粗

        Returns: [{text, is_heading, level, element}, ...]
        """
        import re
        from collections import Counter

        body_sizes = []
        infos = []

        # ── Pass 1: 收集原始信息 ──
        for p in doc.paragraphs:
            text = p.text.strip()
            style_name = (p.style.name if p.style else "").lower().strip()

            sizes = []
            for run in p.runs:
                if run.font.size:
                    sizes.append(run.font.size.pt)
            max_size = max(sizes) if sizes else 0
            is_bold = any(r.bold for r in p.runs if r.bold is not None)

            hs = re.search(r'(heading|标题)\s*(\d+)', style_name)
            heading_style_level = int(hs.group(2)) if hs else None

            infos.append({
                "text": text,
                "style": style_name,
                "max_size": max_size,
                "is_bold": is_bold,
                "char_count": len(text),
                "heading_style_level": heading_style_level,
            })

            # 收集正文字号基准
            if not heading_style_level and len(text) > 30:
                body_sizes.append(max_size)

        body_baseline = Counter(body_sizes).most_common(1)[0][0] if body_sizes else 12

        # ── Pass 2: 判断是否是标题 ──
        for info in infos:
            if not info["text"]:
                info["is_heading"] = False
                continue

            t = info["text"]
            has_heading_style = info["heading_style_level"] is not None

            # 编号模式
            # 编号模式（仅短文本视为标题信号，长文本→列表项）
            has_numbering = info["char_count"] <= 30 and bool(re.match(
                r'^(第[一二三四五六七八九十百千万\d]+[章节篇部]|'
                r'附录|Chapter\s*\d+|'
                r'[一二三四五六七八九十]+[、.]|'
                r'\d+[.、)]|\d+\.\d+|\d+\.\d+\.\d+)',
                t[:20], re.I
            ))

            short = info["char_count"] < 60
            no_end = not any(t.endswith(c) for c in "。！？，；：.!?;,")
            larger = info["max_size"] > body_baseline + 1.5 if info["max_size"] else False

            signals = 0
            if has_heading_style: signals += 2
            if has_numbering: signals += 2
            if short and no_end: signals += 1
            if larger: signals += 1
            if short and info["is_bold"]: signals += 1

            info["is_heading"] = has_heading_style or has_numbering or (short and no_end and larger) or (short and no_end and info["is_bold"] and signals >= 2)

        # ── Pass 2.5: 过滤目录区 ────────────────────────────
        # 策略: 通过"目录"/"Contents"关键词触发，向后扫描并清除连续标题候选
        # 遇到有真实标题样式(heading_style_level)的段落时停止（但跳过触发段落自身）
        for i, info in enumerate(infos):
            t = info["text"]
            if re.match(r'^目\s*[录次]|^Contents?|^Table of Contents', t[:20], re.I):
                for j in range(i, min(i + 30, len(infos))):
                    if j > i and infos[j]["heading_style_level"] is not None:
                        break  # 跳过"目录"自身，后续有标题样式则停止
                    if infos[j]["is_heading"]:
                        infos[j]["is_heading"] = False
                        infos[j]["element"] = "body"
                    elif infos[j]["char_count"] > 50:
                        break

        # ── Pass 3: 分配层级 (H1/H2/H3) ──
        heading_sizes = sorted(set(
            info["max_size"] for info in infos
            if info["is_heading"] and info["max_size"] > 0
        ), reverse=True)

        size_level = {}
        for i, s in enumerate(heading_sizes[:3]):
            size_level[s] = i + 1

        for info in infos:
            if not info["is_heading"]:
                info["level"] = None
                info["element"] = "body"
                continue

            t = info["text"]

            if info["heading_style_level"] is not None:
                info["level"] = min(info["heading_style_level"], 3)
            elif re.match(r'^(第[一二三四五六七八九十百千万\d]+[章节篇部]|附录)', t[:20]):
                info["level"] = 1
            elif re.match(r'^Chapter\s*\d+', t[:15], re.I):
                info["level"] = 1
            elif re.match(r'^[一二三四五六七八九十]+[、.]', t[:8]):
                info["level"] = 2
            elif re.match(r'^\d+\.\d+\.\d+', t[:12]) and len(t) <= 25:
                info["level"] = 3
            elif re.match(r'^\d+\.\d+', t[:8]) and len(t) <= 25:
                info["level"] = 2
            elif re.match(r'^\d+[.、)]', t[:6]) and len(t) <= 20:
                info["level"] = 2
            elif re.match(r'^[（(][一二三四五六七八九十\d]+[）)]', t[:10]) and len(t) <= 20:
                info["level"] = 3
            elif size_level and info["max_size"] in size_level:
                info["level"] = size_level[info["max_size"]]
            else:
                info["level"] = 3

            level_map = {1: "title", 2: "subtitle", 3: "heading"}
            info["element"] = level_map.get(info["level"], "heading")

        self._structure = infos
        return infos

    # ── 格式化核心 ────────────────────────────────────────

    def apply_to_document(self, input_path: str, output_path: str):
        doc = Document(input_path)
        self._doc = doc
        self._apply_page(doc)
        self._detect_structure(doc)  # 先智能分析全文结构
        for idx, p in enumerate(doc.paragraphs):
            self._format(p, idx)
        doc.save(output_path)
        self._doc = None

    def _apply_page(self, doc):
        pg = self.page
        for section in doc.sections:
            if "margin_top" in pg: section.top_margin = Cm(pg["margin_top"])
            if "margin_bottom" in pg: section.bottom_margin = Cm(pg["margin_bottom"])
            if "margin_left" in pg: section.left_margin = Cm(pg["margin_left"])
            if "margin_right" in pg: section.right_margin = Cm(pg["margin_right"])
            self._apply_header_footer(section, pg.get("header", {}))

    def _apply_header_footer(self, section, hdr):
        """应用页眉文字和页脚页码"""
        if not hdr:
            return
        show_header = hdr.get("show_header", False)
        text = hdr.get("text", "")
        show_page_number = hdr.get("show_page_number", False)
        use_chapter = hdr.get("use_chapter_header", False)
        align = hdr.get("page_number_align", "center")
        wda = ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.CENTER)

        if show_header:
            section.header.is_linked_to_previous = False
            hp = section.header.paragraphs[0]
            hp.alignment = wda
            if use_chapter:
                # StyleRef 域代码：自动提取最近的一级标题作为页眉
                hp.clear()
                run = hp.add_run()
                fld_begin = OxmlElement('w:fldChar')
                fld_begin.set(qn('w:fldCharType'), 'begin')
                run._r.append(fld_begin)
                run2 = hp.add_run()
                instr = OxmlElement('w:instrText')
                instr.set(qn('xml:space'), 'preserve')
                instr.text = ' STYLEREF "Heading 1" \\* MERGEFORMAT '
                run2._r.append(instr)
                run3 = hp.add_run()
                fld_end = OxmlElement('w:fldChar')
                fld_end.set(qn('w:fldCharType'), 'end')
                run3._r.append(fld_end)
            elif text:
                hp.text = text

        if show_page_number:
            section.different_first_page_header_footer = False
            footer = section.footer
            footer.is_linked_to_previous = False
            fp = footer.paragraphs[0]
            fp.alignment = wda
            fp.clear()

            run = fp.add_run()
            fld_begin = OxmlElement('w:fldChar')
            fld_begin.set(qn('w:fldCharType'), 'begin')
            run._r.append(fld_begin)

            run2 = fp.add_run()
            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            instr.text = ' PAGE '
            run2._r.append(instr)

            run3 = fp.add_run()
            fld_end = OxmlElement('w:fldChar')
            fld_end.set(qn('w:fldCharType'), 'end')
            run3._r.append(fld_end)

    def _identify_element(self, p, idx: int = None) -> str:
        """识别段落属于哪种文档元素

        优先级: text_structure(前端智能检测) > style_map(用户显式映射) > 智能结构检测 > 传统样式匹配
        """
        import re

        # 0) 前端传入的 text_structure（最高优先级，文本模式由客户端智能检测提供）
        if idx is not None and self.text_structure:
            idx_str = str(idx)
            if idx_str in self.text_structure:
                return self.text_structure[idx_str]

        style_name = (p.style.name if p.style else "").lower().strip()

        # 1) 前端传入的自定义样式映射（最高优先级）
        if self.style_map and style_name in self.style_map:
            return self.style_map[style_name]

        # 2) 智能结构检测（全文分析后的结果）
        if idx is not None and self._structure and idx < len(self._structure):
            detected = self._structure[idx]
            if not detected["text"]:
                return "body"
            if detected["is_heading"]:
                return detected["element"]
            # 非标题但智能结构明确判断为某类元素
            elem = detected.get("element", "body")
            if elem in self.elements or elem == "body":
                return elem

        # 3) 正则兜底：从样式名推断层级
        if "heading" in style_name or "标题" in style_name:
            m = re.search(r'(\d+)', style_name)
            if m:
                level = int(m.group(1))
                if level == 1:
                    return "title"
                elif level == 2:
                    return "subtitle"
                elif level >= 3:
                    return "heading"

        # 4) 正文类
        if style_name in ("normal", "正文", "body text"):
            return "body"

        return "body"

    def _apply_rule(self, run, rule: dict):
        # 跳过图片/绘图运行（不修改绘制元素）
        if run._element.findall(qn('w:drawing')):
            return
        if "font" in rule:
            run.font.name = rule["font"]
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), rule["font"])
        if "size" in rule:
            run.font.size = Pt(self._to_number(rule["size"]))
        if "color" in rule:
            run.font.color.rgb = RGBColor.from_string(rule["color"])
        if "bold" in rule:
            # 将字符串 "true"/"false" 转为布尔值
            bold_val = rule["bold"]
            if isinstance(bold_val, str):
                bold_val = bold_val.lower() == "true"
            run.bold = bold_val

    def _set_outline_level(self, p, level: int):
        """设置段落的大纲级别，0=H1, 1=H2, 2=H3 ...
        确保输出文档在 Word 导航窗格/目录中呈现正确层级"""
        pPr = p._element.get_or_add_pPr()
        for existing in pPr.findall(qn("w:outlineLvl")):
            pPr.remove(existing)
        lvl = OxmlElement("w:outlineLvl")
        lvl.set(qn("w:val"), str(level))
        pPr.append(lvl)

    def _to_number(self, val):
        """将字符串数字转为 float，非数字返回原值"""
        if isinstance(val, (int, float)):
            return val
        try:
            return float(val)
        except (ValueError, TypeError):
            return val

    def _format(self, p, idx: int = None):
        elem = self._identify_element(p, idx)
        rule = self.elements.get(elem) if self.elements else None
        if not rule:
            rule = DEFAULT_ELEMENTS.get(elem, DEFAULT_ELEMENTS["body"])

        pf = p.paragraph_format
        if "line_spacing" in rule: pf.line_spacing = self._to_number(rule["line_spacing"])
        if "space_before" in rule: pf.space_before = Pt(self._to_number(rule["space_before"]))
        if "space_after" in rule: pf.space_after = Pt(self._to_number(rule["space_after"]))
        if "first_line_indent" in rule: pf.first_line_indent = Cm(self._to_number(rule["first_line_indent"]))
        if rule.get("align") in ALIGN_MAP: pf.alignment = ALIGN_MAP[rule["align"]]

        # 设置大纲级别，确保文档结构层级正确
        outline_map = {"title": 0, "subtitle": 1, "heading": 2}
        if elem in outline_map:
            self._set_outline_level(p, outline_map[elem])

        # 如果启用了章节页眉，为标题段落应用 Word 标题样式（供 StyleRef 使用）
        hdr_cfg = self.page.get("header", {})
        if hdr_cfg.get("use_chapter_header", False) and elem in outline_map:
            heading_style_names = ["Heading 1", "Heading 2", "Heading 3"]
            sn = heading_style_names[outline_map[elem]]
            try:
                if hasattr(self, '_doc'):
                    p.style = self._doc.styles[sn]
            except Exception:
                pass

        for run in p.runs:
            self._apply_rule(run, rule)


class DocumentFormatAgent:
    def __init__(self):
        pass

    def format_document(self, target: str, output: str, config: dict):
        if not config:
            return
        formatter = DocumentFormatter(config)
        formatter.apply_to_document(target, output)

    def preview_document(self, target: str, config: dict) -> list:
        """提取文档结构: [{text, style, element, index}, ...]"""
        doc = Document(target)
        formatter = DocumentFormatter(config) if config else None
        if formatter:
            formatter._detect_structure(doc)
        result = []
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if not text:
                continue
            style_name = p.style.name if p.style else "Normal"
            element = formatter._identify_element(p, i) if formatter else "body"
            result.append({
                "index": i,
                "text": text,
                "style": style_name,
                "element": element,
            })
        return result

    def create_doc_from_text(self, text: str, output: str, config: dict):
        """从纯文本创建 docx（每行一个 Normal 段落），然后格式化"""
        doc = Document()
        lines = text.strip().split("\n")
        for line in lines:
            line = line.strip()
            if line:
                p = doc.add_paragraph(line)
                p.style = doc.styles["Normal"]
        doc.save(output)
        if config:
            self.format_document(output, output, config)


def main():
    parser = argparse.ArgumentParser(description="文档格式 Agent v2.0 — 分层格式化引擎")
    parser.add_argument("--stdin", action="store_true", help="从 stdin 读取 JSON")
    parser.add_argument("--preview", action="store_true", help="预览模式：输出文档段落结构")
    parser.add_argument("--text", action="store_true", help="文本模式：从文字创建 docx")
    parser.add_argument("--quiet", "-q", action="store_true", help="静默模式")
    args = parser.parse_args()

    if args.quiet:
        sys.stdout = open(os.devnull, "w", encoding="utf-8", errors="replace")

    if args.stdin:
        raw = sys.stdin.buffer.read()
        cfg = json.loads(raw.decode("utf-8"))
        agent = DocumentFormatAgent()

        if args.preview:
            # 预览模式：返回文档结构 JSON
            target = cfg["target"]
            elements = cfg.get("elements", {})
            page = cfg.get("page", {})
            style_map = cfg.get("style_map", {})
            text_structure = cfg.get("text_structure", {})
            config = {"elements": elements, "page": page, "style_map": style_map, "text_structure": text_structure}
            result = agent.preview_document(target, config)
            # 输出到 stdout（临时恢复以输出 JSON）
            old_stdout = sys.stdout
            if args.quiet:
                sys.stdout = sys.__stdout__
            print(json.dumps(result, ensure_ascii=False))
            if args.quiet:
                sys.stdout = old_stdout

        elif args.text:
            # 文本模式：从文字创建 docx 并格式化
            text = cfg["text"]
            output = cfg["output"]
            elements = cfg.get("elements", {})
            page = cfg.get("page", {})
            style_map = cfg.get("style_map", {})
            text_structure = cfg.get("text_structure", {})
            config = {"elements": elements, "page": page, "style_map": style_map, "text_structure": text_structure}
            agent.create_doc_from_text(text, output, config)

        else:
            # 默认：格式化模式
            target = cfg["target"]
            output = cfg["output"]
            elements = cfg.get("elements", {})
            page = cfg.get("page", {})
            style_map = cfg.get("style_map", {})
            config = {"elements": elements, "page": page, "style_map": style_map}
            agent.format_document(target, output, config)
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
